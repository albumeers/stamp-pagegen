# -*- coding: utf-8 -*-

# Copyright 2026 Jason Drake (jadrake75@gmail.com) and John Lechtanski
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate PF Reference PDFs from stamp album XML files.

Reads combined stamp-album XML input, extracts constant plate flaw and
constant overprint flaw entries, and produces formatted reference documents
with stamp images and descriptions arranged in paginated grids.

The pipeline supports two PDF backends:

    * ``reportlab`` (default): Direct PDF generation via ReportLab.
    * ``word``: python-docx document creation followed by Word COM conversion.

Typical usage from the command line::

    python generate-plateflaws.py --selection ddr --processor reportlab

When ``--selection`` / ``-s`` is omitted, the program prompts interactively.
Use ``--processor`` / ``-p`` to choose ``reportlab`` or ``word``.

Input XML files are read from ``input_directory``; stamp images are resolved
from the sibling ``images`` folder. Output PDFs and timing logs are written
to ``output_directory`` (``baseline_timings.txt`` and ``detailed_timings.txt``).

"""

import logging
import os
import re   ## regex
import sys
import tempfile
import time  # Importing the time module to work with time-related functions

try:
    import pythoncom
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    pythoncom = None
    win32com = None
    HAS_WIN32COM = False

import html
import json
import gc
import io
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    Image = None
    HAS_PIL = False

try:
    import docx
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from docx.shared import Mm
    from docx.shared import Length
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.shared import RGBColor
    HAS_DOCX = True
except ImportError:
    docx = None
    Document = None
    HAS_DOCX = False


# ReportLab imports for integrated PDF backend
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader, simpleSplit
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except Exception:
    # ReportLab may not be installed; the program will fall back to Word backend
    canvas = None
    letter = None
    ImageReader = None
    simpleSplit = None
    pdfmetrics = None
    TTFont = None

# PDF Unit Constants (points per unit)
PTS_PER_MM = 72.0 / 25.4
PTS_PER_INCH = 72.0

# Simple image reader cache to avoid repeated decoding of identical files
_IMAGE_CACHE = {}

def _get_memory_usage_mb():
    """Retrieve current RSS memory usage in MB using psutil or resource fallback."""
    try:
        import psutil
        process = psutil.Process()
        return round(process.memory_info().rss / (1024.0 * 1024.0), 2)
    except Exception:
        pass
    try:
        import resource
        rusage = resource.getrusage(resource.RUSAGE_SELF)
        return round(rusage.ru_maxrss / 1024.0, 2)
    except Exception:
        pass
    return 0.0


def _resolve_config_params(cli_args=None, font_mappings=None):
    """Resolves configuration parameters adhering strictly to a 3-tier precedence hierarchy:
    
    1. Tier 1: Calculated / Hard-coded Defaults
    2. Tier 2: mapping.json overrides
    3. Tier 3: Command-Line Arguments (argparse) overrides
    
    Returns:
        dict: Resolved configuration parameters.
    """
    calc_input_dir = _get_documents_dir()
    config = {
        'selection': None,
        'processor': 'reportlab',
        'input-dir': calc_input_dir,
        'output-dir': tempfile.gettempdir(),
        'images-dir': os.path.join(calc_input_dir, 'images'),
        'image-type': 'jpeg',
        'image-quality': 85,
        'font-mappings': None
    }

    # Step 1: Load mapping.json into mapping_data (Tier 2)
    mapping_data = font_mappings
    if mapping_data is None:
        json_path = Path('mapping.json')
        if not json_path.exists():
            json_path = Path(__file__).parent / 'mapping.json'
        if json_path.exists():
            try:
                with open(json_path, 'r', encoding='utf-8') as f:
                    mapping_data = json.load(f)
            except Exception as e:
                print(f"Warning: Failed to load mapping from {json_path}: {e}")

    # Extract font-mappings and Tier 2 parameter overrides
    if isinstance(mapping_data, dict):
        if 'font-mappings' in mapping_data:
            config['font-mappings'] = mapping_data['font-mappings']
        else:
            config['font-mappings'] = mapping_data

        if 'input-dir' in mapping_data and mapping_data['input-dir']:
            config['input-dir'] = str(mapping_data['input-dir'])
            config['images-dir'] = os.path.join(config['input-dir'], 'images')

        if 'output-dir' in mapping_data and mapping_data['output-dir']:
            config['output-dir'] = str(mapping_data['output-dir'])

        if 'images-dir' in mapping_data and mapping_data['images-dir']:
            config['images-dir'] = str(mapping_data['images-dir'])

        if 'processor' in mapping_data and mapping_data['processor']:
            config['processor'] = str(mapping_data['processor']).strip().lower()

        if 'image-type' in mapping_data and mapping_data['image-type']:
            val = str(mapping_data['image-type']).strip().lower()
            if val in ('png', 'jpeg', 'jpg'):
                config['image-type'] = 'png' if val == 'png' else 'jpeg'

        if 'image-quality' in mapping_data and mapping_data['image-quality'] is not None:
            try:
                q = int(mapping_data['image-quality'])
                if 1 <= q <= 100:
                    config['image-quality'] = q
            except (ValueError, TypeError):
                pass

        if 'selection' in mapping_data and mapping_data['selection']:
            config['selection'] = str(mapping_data['selection']).strip().lower()

    elif hasattr(mapping_data, 'font_mappings'):
        config['font-mappings'] = getattr(mapping_data, 'font_mappings')
    elif mapping_data is not None:
        config['font-mappings'] = mapping_data

    # Step 2: Merge Tier 3 (Command-Line Arguments overrides)
    if cli_args:
        if getattr(cli_args, 'selection', None):
            config['selection'] = str(cli_args.selection).strip().lower()

        cli_proc = getattr(cli_args, 'processor', None)
        if cli_proc:
            config['processor'] = str(cli_proc).strip().lower()

        if getattr(cli_args, 'input_dir', None):
            config['input-dir'] = str(cli_args.input_dir)
            config['images-dir'] = os.path.join(config['input-dir'], 'images')

        if getattr(cli_args, 'output_dir', None):
            config['output-dir'] = str(cli_args.output_dir)

        if getattr(cli_args, 'images_dir', None):
            config['images-dir'] = str(cli_args.images_dir)

        if getattr(cli_args, 'image_type', None):
            val = str(cli_args.image_type).strip().lower()
            if val in ('png', 'jpeg', 'jpg'):
                config['image-type'] = 'png' if val == 'png' else 'jpeg'

        if getattr(cli_args, 'image_quality', None) is not None:
            try:
                q = int(cli_args.image_quality)
                if 1 <= q <= 100:
                    config['image-quality'] = q
            except (ValueError, TypeError):
                pass

    return config


def _load_mapping_config(font_mappings=None):
    """Load configuration options adhering to the 3-tier precedence hierarchy.

    Args:
        font_mappings: Optional dict, list, or object containing configuration.

    Returns:
        dict: A dict containing resolved configuration options.
    """
    return _resolve_config_params(font_mappings=font_mappings)


def _register_fonts(font_mappings=None):
    """Register ReportLab TrueType fonts used by the PDF backend.

    Priority:
    1. If ``font_mappings`` argument is provided (or contains a ``'font-mappings'`` key),
       use those mappings.
    2. If ``font_mappings`` is None/omitted, attempt to read from a local ``mapping.json`` file.

    Args:
        font_mappings: Optional dict, list, or object containing font mapping definitions.

    Side Effects:
        Registers mapped fonts with ``pdfmetrics`` and prints the registered fonts.
    """
    if pdfmetrics is None:
        return
    registered = {}
    config = _load_mapping_config(font_mappings)
    resolved_mappings = config['font-mappings']

    # 3. Register resolved font mappings with ReportLab
    if resolved_mappings:
        mapping_dict = {}
        if isinstance(resolved_mappings, dict):
            mapping_dict = resolved_mappings
        elif isinstance(resolved_mappings, (list, tuple)):
            for item in resolved_mappings:
                if isinstance(item, dict):
                    name = item.get('name') or item.get('fontName') or item.get('alias') or item.get('family')
                    path = item.get('path') or item.get('file') or item.get('filename')
                    if name and path:
                        mapping_dict[name] = path
                elif isinstance(item, (list, tuple)) and len(item) >= 2:
                    mapping_dict[item[0]] = item[1]

        for font_name, font_path in mapping_dict.items():
            names_to_register = {font_name}
            if font_name.lower() == 'castletlig':
                names_to_register.add('CastleTLig')
            
            expanded_path = Path(os.path.expandvars(os.path.expanduser(str(font_path))))
            actual_path = expanded_path if expanded_path.exists() else None

            # Fallback to C:\Windows\Fonts if file isn't found at configured path on Windows
            if not actual_path and sys.platform == 'win32':
                win_font = Path("C:/Windows/Fonts") / expanded_path.name
                if win_font.exists():
                    actual_path = win_font

            if actual_path:
                for name in names_to_register:
                    if name not in pdfmetrics.getRegisteredFontNames():
                        try:
                            pdfmetrics.registerFont(TTFont(name, str(actual_path)))
                            registered[name] = str(actual_path)
                        except Exception as e:
                            print(f"Warning: Could not register font {name} from {actual_path}: {e}")
            else:
                print(f"Warning: Font file for '{font_name}' not found at '{expanded_path}'")

    if registered:
        try:
            print(f"ReportLab: registered fonts: {registered}")
        except Exception:
            pass


def _resolve_image_path(images_directory, image_field):
    """Resolve an absolute path to a stamp image file.

    Args:
        images_directory: Directory containing the album ``images`` folder.
        image_field: Relative image path from XML data, e.g.
            ``images/Germany - DDR/OD1 II.png``.

    Returns:
        pathlib.Path: Absolute, resolved path to the image file.
    """
    base = Path(images_directory).parent
    return (base / image_field).resolve()


def generate_pdf(data_list, output_pdf_path, images_directory, label_print=False, page_title='', page_description='', font_mappings=None):
    """Generate a PF reference PDF using ReportLab.

    Builds paginated grid layouts that mirror the python-docx printing
    logic: stamp images with captions, page titles, and optional label mode.

    Args:
        data_list: Parsed plate-fault data rows to render.
        output_pdf_path: Destination path for the generated PDF file.
        images_directory: Base directory used to resolve image paths.
        label_print: If True, use a 4x10 label layout; otherwise 5x6.
        page_title: Title printed at the top of each page.
        page_description: Subtitle printed below the page title.
        font_mappings: Optional font mapping definitions or dict.

    Raises:
        RuntimeError: If ReportLab is not installed or importable.
    """
    if canvas is None:
        raise RuntimeError('ReportLab not available')

    # Register fonts once
    _register_fonts(font_mappings)

    # Layout
    if label_print:
        cols = 4
        rows = 10
    else:
        cols = 5
        rows = 6

    page_size = letter if (isinstance(letter, (list, tuple)) and len(letter) == 2) else (612.0, 792.0)
    top_margin = 18 * PTS_PER_MM
    bottom_margin = 18 * PTS_PER_MM
    left_margin = 18 * PTS_PER_MM
    right_margin = 18 * PTS_PER_MM
    header_h = 26 * PTS_PER_MM
    caption_h = 14 * PTS_PER_MM - (3 * PTS_PER_MM)
    pad = 0 * PTS_PER_MM
    caption_pad_top = 2 * PTS_PER_MM
    caption_pad_bottom = 0 * PTS_PER_MM
    EXTRA_IMAGE_HEIGHT = 2 * PTS_PER_MM
    CELLS_OFFSET_X = 3 * PTS_PER_MM
    CELLS_OFFSET_Y = 8 * PTS_PER_MM
    IMAGE_SCALE_FACTOR = 4.5
    target_img_size = 1.2 * PTS_PER_INCH * IMAGE_SCALE_FACTOR

    # Normalize input data to sections list: [('Constant Plate Flaws', list1), ('Constant Overprint Flaws', list2)]
    if isinstance(data_list, list) and data_list and isinstance(data_list[0], tuple) and len(data_list[0]) == 2 and isinstance(data_list[0][1], list):
        sec_list = data_list
    else:
        sec_list = [(page_title or 'Constant Plate Flaws', data_list)]

    os.makedirs(os.path.dirname(os.path.abspath(str(output_pdf_path))), exist_ok=True)
    c = canvas.Canvas(str(output_pdf_path), pagesize=page_size)
    page_w, page_h = page_size
    usable_w = page_w - left_margin - right_margin
    usable_h = page_h - top_margin - bottom_margin
    grid_h = usable_h - header_h
    cell_w = usable_w / cols
    cell_h = (grid_h / rows) + EXTRA_IMAGE_HEIGHT

    # warn if target image too big
    try:
        if target_img_size > (cell_w - 2 * pad) or target_img_size > (cell_h - caption_h - 2 * pad):
            print(f"ReportLab: target image size ({target_img_size:.2f} pts) exceeds cell box size (w={cell_w - 2*pad:.2f}, h={cell_h - caption_h - 2*pad:.2f}).")
    except Exception:
        pass

    registered_fonts = set(pdfmetrics.getRegisteredFontNames()) if pdfmetrics else set()
    title_font = 'CastleTLig' if 'CastleTLig' in registered_fonts or 'CastleTLig' in registered_fonts else 'Helvetica-Bold'
    body_font = 'Verdana' if 'Verdana' in registered_fonts else 'Helvetica'
    subtitle_font = 'CastleTLig' if 'CastleTLig' in registered_fonts or 'CastleTLig' in registered_fonts else body_font

    # Load image configuration options (image-type and image-quality)
    mapping_config = _load_mapping_config(font_mappings)
    cfg_image_type = mapping_config.get('image-type', 'jpeg')
    cfg_image_quality = mapping_config.get('image-quality', 85)

    def _process_single_image_file(path_obj, max_dim=350, img_type=cfg_image_type, quality=cfg_image_quality):
        if not path_obj or not path_obj.exists():
            return None
        try:
            with Image.open(str(path_obj)) as im:
                orig_w, orig_h = im.size
                if max(orig_w, orig_h) > max_dim:
                    scale = max_dim / float(max(orig_w, orig_h))
                    new_w = max(1, int(orig_w * scale))
                    new_h = max(1, int(orig_h * scale))
                    resized = im.resize((new_w, new_h), Image.Resampling.LANCZOS)
                else:
                    resized = im.copy()

                buf = io.BytesIO()
                if resized.mode in ('RGBA', 'LA') or (resized.mode == 'P' and 'transparency' in resized.info):
                    resized.save(buf, format='PNG', optimize=True)
                elif img_type == 'png':
                    resized.save(buf, format='PNG', optimize=True)
                else:
                    resized.convert('RGB').save(buf, format='JPEG', quality=quality, optimize=True)
                buf.seek(0)
                img_reader = ImageReader(buf) if ImageReader else None
                return (img_reader, orig_w, orig_h)
        except Exception:
            try:
                img_reader = ImageReader(str(path_obj)) if ImageReader else None
                return (img_reader, None, None)
            except Exception:
                return None

    def _preload_page_images(page_rows):
        page_paths = {}
        for row_list in page_rows:
            for item in row_list:
                if len(item) > 1 and item[1]:
                    p_obj = _resolve_image_path(images_directory, item[1])
                    k = str(p_obj)
                    if k not in _IMAGE_CACHE:
                        page_paths[k] = p_obj
        if page_paths:
            with ThreadPoolExecutor() as executor:
                future_to_key = {executor.submit(_process_single_image_file, p_obj): k for k, p_obj in page_paths.items()}
                for future in future_to_key:
                    k = future_to_key[future]
                    try:
                        res = future.result()
                        if res:
                            _IMAGE_CACHE[k] = res
                    except Exception:
                        pass

    left_inset = 0.75 * PTS_PER_INCH
    other_inset = 0.5 * PTS_PER_INCH
    border_x = left_inset
    border_y = other_inset
    border_w = page_w - left_inset - other_inset
    border_h = page_h - other_inset - other_inset

    first_section = True
    for sec_title, sec_data in sec_list:
        if not sec_data:
            continue

        if not first_section:
            # Render visual Divider Page between sections
            c.setLineWidth(1.5)
            c.setStrokeColorRGB(0.0, 0.0, 0.0)
            c.rect(border_x, border_y, border_w, border_h, stroke=1, fill=0)

            center_y = border_y + (border_h / 2.0)

            div_main = (page_title or sec_title).upper()
            c.setFont(title_font, 28)
            c.drawCentredString(page_w / 2.0, center_y + (12 * PTS_PER_MM), div_main)

            if page_description and page_description.strip():
                c.setFont(subtitle_font, 18)
                c.drawCentredString(page_w / 2.0, center_y + (2 * PTS_PER_MM), page_description.upper())

            c.setFont(subtitle_font, 14)
            c.drawCentredString(page_w / 2.0, center_y - (8 * PTS_PER_MM), sec_title.upper())

            c.showPage()

        first_section = False

        temp_list = sec_data
        list2 = [temp_list[i:i+cols] for i in range(0, len(temp_list), cols)]
        list3 = [list2[i:i+rows] for i in range(0, len(list2), rows)]

        image_last_page_map = {}
        for p_idx, page in enumerate(list3):
            for row_list in page:
                for item in row_list:
                    if len(item) > 1 and item[1]:
                        img_path = _resolve_image_path(images_directory, item[1])
                        image_last_page_map[str(img_path)] = p_idx

        for page_idx, page in enumerate(list3):
            # Pre-load images required for the current page
            _preload_page_images(page)
            # outer border
            try:
                c.setLineWidth(1.5)
                c.setStrokeColorRGB(0.0, 0.0, 0.0)
                c.rect(border_x, border_y, border_w, border_h, stroke=1, fill=0)
            except Exception:
                pass

            c.setFont(title_font, 26)
            content_top = border_y + border_h
            title_y = content_top - (12 * PTS_PER_MM)
            subtitle_y = title_y - (7 * PTS_PER_MM)
            main_title = (page_title or sec_title or 'PF Reference').upper()
            c.drawCentredString(page_w / 2.0, title_y, main_title)

            c.setFont(subtitle_font, 13)
            if page_description and page_description.strip():
                sub_text = page_description.upper()
            elif label_print:
                sub_text = "CONSTANT PLATE FLAW LABELS"
            else:
                sub_text = sec_title.upper()
            c.drawCentredString(page_w / 2.0, subtitle_y, sub_text)

            for r, row_list in enumerate(page):
                for col_idx, item in enumerate(row_list):
                    x0 = left_margin + CELLS_OFFSET_X + col_idx * cell_w
                    y_top = page_h - top_margin - header_h - (r * cell_h) + CELLS_OFFSET_Y
                    y0 = y_top - cell_h

                    box_w = cell_w
                    box_h = cell_h - caption_h
                    img_box_size = min(box_w, box_h)

                    # image box coords
                    img_box_w = cell_w
                    img_box_h = img_box_size
                    img_box_x = x0
                    img_box_y = y0 + caption_h + (box_h - img_box_h) / 2.0

                    # draw image
                    try:
                        img_field = item[1]
                        img_path = _resolve_image_path(images_directory, img_field)
                        key = str(img_path)
                        img_entry = _IMAGE_CACHE.get(key)
                        if not img_entry and img_path.exists() and ImageReader:
                            img_entry = (ImageReader(str(img_path)), None, None)
                        if img_entry and img_entry[0]:
                            img_reader, orig_w, orig_h = img_entry
                            iw, ih = (orig_w, orig_h) if (orig_w and orig_h) else img_reader.getSize()
                            draw_w = min(target_img_size, img_box_w)
                            draw_h = min(target_img_size, img_box_h)
                            if iw and ih:
                                ratio = iw / ih
                                if draw_w / draw_h > ratio:
                                    draw_w = draw_h * ratio
                                else:
                                    draw_h = draw_w / ratio
                            draw_x = img_box_x + (img_box_w - draw_w) / 2.0
                            draw_y = img_box_y + (img_box_h - draw_h) / 2.0
                            c.drawImage(img_reader, draw_x, draw_y, width=draw_w, height=draw_h, preserveAspectRatio=True, anchor='c')
                    except Exception:
                        pass

                    # caption
                    try:
                        raw_caption = item[9] if len(item) > 9 else Path(item[1]).name
                    except Exception:
                        raw_caption = Path(item[1]).name
                    caption_text = html.unescape(raw_caption.replace('\\n', '\n')).lstrip()
                    caption_font = body_font
                    caption_size = 4.5
                    leading = 1.0
                    max_width = cell_w
                    lines = []
                    for para in caption_text.splitlines():
                        if para.strip() == '':
                            lines.append('')
                        else:
                            wrapped = simpleSplit(para, caption_font, caption_size, max_width)
                            if wrapped:
                                lines.extend(wrapped)
                            else:
                                lines.append(para)
                    available_caption_height = caption_h - caption_pad_top - caption_pad_bottom
                    max_lines = int(available_caption_height // (caption_size + leading))
                    if max_lines < 1:
                        max_lines = 1
                    if len(lines) > max_lines:
                        lines = lines[:max_lines]
                        if len(lines[-1]) > 3:
                            lines[-1] = lines[-1][:-3] + '...'
                    caption_area_top = y0 + caption_h - caption_pad_top
                    y_first = caption_area_top - caption_size
                    for i, ln in enumerate(lines):
                        y_pos = y_first - i * (caption_size + leading)
                        c.setFont(caption_font, caption_size)
                        c.drawCentredString(x0 + cell_w / 2.0, y_pos, ln)

            c.showPage()

            # Evict images from cache whose last required page has been rendered
            keys_to_evict = [k for k, last_idx in image_last_page_map.items() if last_idx <= page_idx and k in _IMAGE_CACHE]
            if keys_to_evict:
                for k in keys_to_evict:
                    del _IMAGE_CACHE[k]
                gc.collect()

    c.save()
    print(f"Generated PDF: {os.path.abspath(str(output_pdf_path))}")
    _IMAGE_CACHE.clear()
    gc.collect()


# Program Control Flags
label_print = False

def _get_documents_dir():
    """Resolves the user's Documents directory via Windows SHGetFolderPathW or Path.home() fallback."""
    if sys.platform == 'win32':
        try:
            import ctypes.wintypes
            CSIDL_PERSONAL = 5
            SHGFP_TYPE_CURRENT = 0
            buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
            ctypes.windll.shell32.SHGetFolderPathW(None, CSIDL_PERSONAL, None, SHGFP_TYPE_CURRENT, buf)
            if buf.value:
                return str(Path(buf.value))
        except Exception:
            pass
    return str(Path.home() / 'Documents')

# Define input, output, and image directories (environment-aware with relative fallback)
default_input_dir = _get_documents_dir()
default_output_dir = tempfile.gettempdir()

input_directory = os.environ.get('STAMP_XML_DIR', default_input_dir)
output_directory = os.environ.get('STAMP_OUTPUT_DIR', default_output_dir)
images_directory = os.path.join(input_directory, 'images')

## position of data items within 'data' sublists within 'data_list'

#      i_data = 0
global i_image ; i_image = 1 
#      i_dimensions = 2
global i_denomination ; i_denomination = 3 
global i_color ; i_color = 4
global i_cell_desc ; i_cell_desc = 5
#      i_cat_nums = 6 catalog numbers
#      i_cat1 = 7 Scott
global i_cat2 ; i_cat2 = 8 ## Michel
global i_rdesc ; i_rdesc = 9 ## description taken from '<row-set>'
global i_date ; i_date = 10 ## year taken from '<set issue='
global i_sort ; i_sort = 11
#      i_sort_special = 0 ## special use: C = airmail, O = official, etc. position 0, length 1
global i_sort_year ; i_sort_year = 1 ## year: position 1, length 4
#      i_sort_cat_num = 5 ## Michel number: justified right, padded with zeroes: position 5, length 10
global i_sort_field ; i_sort_field = 15 ## Michel plate fault Roman number or field: position 15, variable length
global start_time ; start_time = time.time()

# Initialize the COM library if available
if HAS_WIN32COM and pythoncom is not None:
    try:
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch('Word.Application')
    except Exception:
        word = None
else:
    word = None

def _extract_section_data_list(raw_xml_list, section_title):
    raw_subset = capture_raw_constant_plate_flaw_data(raw_xml_list, target_title=section_title)
    if not raw_subset:
        return []
    d_list = create_data_list(raw_subset)
    d_list = extract_selectively(d_list)
    d_list = append_row_set_description(d_list)
    d_list = removeDenominationColor(d_list)
    d_list = remove_duplicates_from_data_list(d_list, i_image, i_rdesc)
    return d_list


def createPFRefAlbumPages(input_file_list, output_file, image_directory_location, processor='reportlab'):
    """Parses album XML datasets and creates PDF reference documents.

    Args:
        input_file_list (list): List of XML dataset file paths to combine.
        output_file (str): Base path for generated reference PDF/DOCX.
        image_directory_location (str): Directory containing stamp images.
        processor (str): PDF engine/processor to use: ``'reportlab'`` or ``'word'``.

    Returns:
        tuple[list, list]: Pair of lists containing parsed dataset rows when
        found.
    """
    # Add timing/profiling hooks to capture durations for major phases
    timings = {}
    start_total = time.perf_counter()
    # record requested processor for diagnostics
    timings['processor_requested'] = str(processor)
    print(f"Processor requested: {processor}")

    # 1. Combine input XML files
    t0 = time.perf_counter()
    combined_xml = combine_input_xml_files(input_file_list)
    timings['combine_input_xml_files'] = time.perf_counter() - t0

    # 2. Split into raw XML list
    t0 = time.perf_counter()
    raw_xml_list = splitCombinedXML(combined_xml)
    del combined_xml
    timings['splitCombinedXML'] = time.perf_counter() - t0

    # 3. Extract sections for Constant Plate Flaws and Constant Overprint Flaws
    t0 = time.perf_counter()
    sections = []

    pf_list = _extract_section_data_list(raw_xml_list, "Constant Plate Flaws")
    if pf_list:
        sections.append(("Constant Plate Flaws", pf_list))

    op_list = _extract_section_data_list(raw_xml_list, "Constant Overprint Flaws")
    if op_list:
        sections.append(("Constant Overprint Flaws", op_list))

    del raw_xml_list
    timings['process_sections'] = time.perf_counter() - t0

    # Exit early if no matching data across any section
    if not sections:
        timings['total'] = time.perf_counter() - start_total
        _write_timings_log(timings, note='no_data_found')
        return

    # 4. Optionally generate PDF directly using integrated ReportLab (faster)
    if processor and processor.lower() == 'reportlab':
        t0 = time.perf_counter()
        try:
            # call integrated generator with consolidated sections
            generate_pdf(sections, output_file, image_directory_location, label_print, page_title if 'page_title' in globals() else '', page_description if 'page_description' in globals() else '')
            timings['reportlab_generate_pdf'] = time.perf_counter() - t0
            timings['total'] = time.perf_counter() - start_total
            _write_timings_log(timings)
            del sections
            gc.collect()
            return
        except Exception as e:
            timings['reportlab_generate_pdf_error'] = time.perf_counter() - t0
            print(f'ERROR — ReportLab PDF generation failed: {e}\nFalling back to Word pipeline')

    # 5. Create the in-memory .docx document (fallback or default)
    t0 = time.perf_counter()
    docum = createPFDocx(sections)
    timings['createPFDocx'] = time.perf_counter() - t0

    # 6. Write the .docx file to disk
    t0 = time.perf_counter()
    writeDocxPfRefDoc(docum)
    timings['writeDocxPfRefDoc'] = time.perf_counter() - t0

    # 7. Convert .docx to PDF via Word COM (slow step)
    t0 = time.perf_counter()
    writePdfPfRefDoc()
    timings['writePdfPfRefDoc'] = time.perf_counter() - t0

    # Finalize total time and write timings to log
    timings['total'] = time.perf_counter() - start_total
    _write_timings_log(timings)
    
def append_row_set_description(data_list) :
    """Attach row-set descriptions and issue years to data rows.

    Walks ``data_list`` entries of type ``row_set_description``, ``set``, and
    ``data``, pairing each data row with the most recent description and issue
    year. Skips data rows whose image field contains ``'label'``.

    Args:
        data_list: Parsed album data containing row-set, set, and data entries.

    Returns:
        list: Filtered data rows with row-set description and issue year
        appended as additional fields.
    """
    temp = []
    for i in data_list :
        if i[0] == 'row_set_description' : ## get plate fault description
            rsdesc = i[1]
        if i[0] == 'set' and i[1][0:4].isnumeric() : ## get year   
            issue = i[1]
        if i[0] == 'data' and 'label' not in i[i_image] : ## append them to 'data' list
            i = i + [rsdesc] 
            i = i + [issue]
            temp.append(i)
            
    # Examples
    # ['data', 'images/Germany - DDR/OD1 II.png', '18x18', '', '', '', '', '', '', 'Large Blank Area in\\nEmblem at Right\\n(Position Unknown)\\n(OD2 PF II)', '1965']
    # ['data', 'images/Germany - DDR/OD3 I.png', '18x18', '', '', '', '', '', '', 'Burr on Inside of the\\nFirst &quot;D&quot; of &quot;DDR&quot;\\n(Position 6)\\n(OD3 PF I)', '1965-66']
    # ['data', 'images/Germany - DDR/OD3 II.png', '18x18', '', '', '', '', '', '', 'Outside Right of Leaves\\nis Broken in Emblem\\n(Position 6)\\n(OD3 PF II)', '1965-66']
            
    return temp 


def capture_raw_constant_plate_flaw_data(raw_xml_list, target_title=None):
    """Extract XML lines for the specified constant plate/overprint section.

    Args:
        raw_xml_list: List of XML lines from ``splitCombinedXML()``.
        target_title: Optional issue title to match (e.g. ``'Constant Plate Flaws'``).

    Returns:
        list: Filtered XML lines for the target section.
    """
    if target_title is None:
        try:
            target_title = constant_title
        except NameError:
            target_title = "Constant Plate Flaws"

    file = []
    still_on_PF_page = False
    for i in raw_xml_list:
        if '<page ' in i:
            processPage(i)
        if f'<set issue="{target_title}' in i:
            still_on_PF_page = True
        if '</page>' in i:
            still_on_PF_page = False
        if still_on_PF_page:
            file.append(i)
    return file

def cleanUpPFDesc(string):
    """Normalize a plate-fault description string for printing.

    Converts HTML entity quotes and literal ``\\n`` sequences to their
    printable characters, then removes leading whitespace.

    Args:
        string: Raw description text from XML or data rows.

    Returns:
        str: Cleaned description ready for DOCX or PDF output.
    """
    # Replace HTML encoded quotes and literal \n with actual characters
    string = string.replace('&quot;', '"').replace('\\n', '\n')

    # Remove leading whitespace characters, including newlines
    string = string.lstrip()

    return string

def cleanUpPathName(string):
    """Normalize a file path string for Windows-style use.

    Strips surrounding whitespace, converts forward slashes to backslashes,
    and removes trailing backslashes.

    Args:
        string: Path string to normalize.

    Returns:
        str: Normalized path string.
    """
    string = string.strip().replace('/', '\\\\')
    if string[-2:] == '\\\\':
        string = string[:-2]
    if string.endswith('\\'):
        string = string[:-1]
    return string

_XML_FILE_CACHE = {}

def combine_input_xml_files(input_file_list):
    chunks = []
    for single_file in input_file_list:
        if single_file not in _XML_FILE_CACHE:
            with open(single_file, 'r', encoding='utf-8') as handle:
                _XML_FILE_CACHE[single_file] = handle.read()
        chunks.append(_XML_FILE_CACHE[single_file])
    return ''.join(chunks)

def create_data_list(raw_xml_list):
    """Parse raw XML lines into structured album data rows.

    Identifies each XML element type and dispatches to the appropriate
    ``process*()`` handler. Collects page metadata, row-set descriptions,
    set issues, and stamp ``data`` rows into a single list.

    Args:
        raw_xml_list: List of XML lines from ``splitCombinedXML()``.

    Returns:
        list: Parsed data entries, each a list whose first element is a
        record type such as ``'set'``, ``'row_set_description'``, or
        ``'data'``.
    """

    ## parse, identify and process each line of XML
    
    data_list = []
    
    for line in raw_xml_list:
        if line[0:4] == '<!--':
            processComment(line)
        elif line[0:7] == '<album>':
            processAlbum(line)
        elif line[0:12] == '<column-set ':
            processColumnSet(line)
        elif line[0:12] == '<column-set>':
            processColumnSet(line)
        elif line[0:10] == '<comp-set ':
            processCompSet(line)
        elif line[0:10] == '<comp-set>':
            processCompSet(line)
        elif line[0:15] == '<content-items>':
            processCompSet(line)
            processContentItems(line)
        elif '!DOCTYPE' in line:
            processDoctype(line)
        elif line[0:6] == '<item>':
            processItem(line)
        elif line[0:6] == '<page ':
            res = processPage(line)
            for i in range( 0, len(res) ) : data_list.append(res[i])
        elif line[0:9] == '<row-set ':
            res = processRowSet(line)
            data_list.append(res)
        elif line[0:9] == '<row-set>':
            processRowSet(line)
        elif line[0:3] == '<s ':
            res = processS(line)
            data_list.append(res)
        elif line[0:3] == '<s>':
            res = processS(line)
            data_list.append(res)
        elif line[0:4] == '<set':
            res = processSet(line)
            data_list.append(res)
        elif line[0:12] == '<title-page ':
             processTitlePage(line)       
        elif line[0:8] == '</album>':
            processZAlbum(line)
        elif line[0:13] == '</column-set>':
            processZColumnSet(line)
        elif line[0:11] == '</comp-set>':
            processZCompSet(line)
        elif line[0:16] == '</content-items>':
            processZContentItems(line)
        elif line[0:7] == '</item>':
            processZItem(line)
        elif line[0:7] == '</page>':
            processZPage(line)
        elif line[0:10] == '</row-set>':
            processZRowSet(line)
        elif line[0:4] == '</s>':
            processZS(line)
        elif line[0:6] == '</set>':
            processZSet(line)
        elif line[0:13] == '</set-tenant>' :
            processZSetTenant(line)
        elif line[0:13] == '</title-page>' :
            processZTitlePage(line)
        else:
            1==1
            #print('\nERROR: unknown XML element:\n', line)    

    # Examples:
    # ['set', 'Constant Plate Flaws']
    # ['set', '1965']
    # ['row_set_description', 'Outer Line of Emblem\\nis Broken at Top\\n(Position Unknown)\\n(OD1 PF I)']
    # ['data', 'images/Germany - DDR/OD1 I.png', '18x18', '', '', '', '', '', '']
    # ['row_set_description', 'Large Blank Area in\\nEmblem at Right\\n(Position Unknown)\\n(OD1 PF II)']
    # ['data', 'images/Germany - DDR/OD1 II.png', '18x18', '', '', '', '', '', '']
    # ['row_set_description', 'OB1 PF I']

    return data_list

def createPFDocx(data_list):
    """Build an in-memory PF reference Word document."""
    if isinstance(data_list, list) and data_list and isinstance(data_list[0], tuple) and len(data_list[0]) == 2 and isinstance(data_list[0][1], list):
        sec_list = data_list
    else:
        sec_list = [('Constant Plate Flaws', data_list)]

    first_page = True
    doc1 = Document()

    sections = doc1.sections
    section = sections[0]
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(22)
    section.right_margin = Mm(12)

    sec_pr = doc1.sections[0]._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'text')
    for border_name in ('top', 'bottom', 'left', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '8')
        border_el.set(qn('w:space'), '0')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el)
    if not label_print:
        sec_pr.append(pg_borders)

    for idx, (sec_title, sec_data) in enumerate(sec_list):
        if not sec_data:
            continue
        if idx > 0:
            doc1.add_page_break()
            div_main = (page_title or sec_title).upper()
            para1 = doc1.add_paragraph()
            format1 = para1.paragraph_format
            format1.space_before = Pt(120)
            format1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = para1.add_run(div_main)
            run1.font.name = 'CastleTLig'
            run1.font.size = Pt(28)

            if page_description and page_description.strip():
                para2 = doc1.add_paragraph()
                format2 = para2.paragraph_format
                format2.space_before = Pt(10)
                format2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run2 = para2.add_run(page_description.upper())
                run2.font.name = 'CastleTLig'
                run2.font.size = Pt(18)

            para3 = doc1.add_paragraph()
            format3 = para3.paragraph_format
            format3.space_before = Pt(15)
            format3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = para3.add_run(sec_title.upper())
            run3.font.name = 'CastleTLig'
            run3.font.size = Pt(14)

            doc1.add_page_break()
            first_page = False

        printDocxPages(doc1, first_page, sec_data, 'all', title=sec_title)
        first_page = False

    return doc1


def create_file_list(directory, filenames):
    """Build full paths from a directory and filename list.

    Args:
        directory: Base directory for the files.
        filenames: Bare filenames to join with ``directory``.

    Returns:
        list[str]: Full paths, one per filename, in the same order.
    """
    return [os.path.join(directory, filename) for filename in filenames]

def extract_selectively(data_list):
    """Keep only data, row-set description, and set entries.

    Filters ``data_list`` to entries whose type is ``'data'`` (with a
    non-empty image field), ``'row_set_description'``, or ``'set'``.

    Args:
        data_list: Full parsed album data from ``create_data_list()``.

    Returns:
        list: Filtered data entries used by later pipeline steps.
    """
    temp = []
    for i in data_list :
        if i[0] == 'data' and i[1] != '' :
            temp.append(i)
        if i[0] == 'row_set_description' :
            temp.append(i)
        if i[0] == 'set' :
            temp.append(i)
                
    # Examples:
    # ['set', 'Constant Plate Flaws']
    # ['set', '1965-66']
    # ['row_set_description', 'Burr on Inside of the\\nFirst &quot;D&quot; of &quot;DDR&quot;\\n(Position 6)\\n(OD3 PF I)']
    # ['data', 'images/Germany - DDR/OD3 I.png', '18x18', '', '', '', '', '', '']
                
    return temp

def paragraph_format_run(cell):
    """Prepare a centered table cell paragraph for content insertion.

    Uses the cell's first paragraph, applies zero spacing and centered
    alignment, and adds an empty run for images or text.

    Args:
        cell: ``python-docx`` table cell to format.

    Returns:
        tuple: ``(paragraph, paragraph_format, run)`` for further editing.
    """
    paragraph = cell.paragraphs[0]
    format = paragraph.paragraph_format
    run = paragraph.add_run()
    
    format.space_before = Pt(0)
    format.space_after = Pt(0)
    format.line_spacing = 1.0
    format.alignment = WD_ALIGN_PARAGRAPH.CENTER
         
    return paragraph, format, run

def printHeadings(doc, reference_year, title=None):
    """Add page title and subtitle paragraphs to a DOCX document.

    Uses global ``page_title``, ``page_description``, ``label_print``, and
    title to build centered CastleTLig headings at the top of each page.

    Args:
        doc: ``python-docx`` document being built.
        reference_year: Reference year label (currently unused).
        title: Optional section title.
    """
    global first_page
    if title is None:
        try:
            title = constant_title
        except NameError:
            title = "Constant Plate Flaws"

    main_title = (page_title or title or '').upper()

    ## ----- Page Heading

    # Creating paragraph
    para = doc.add_paragraph()

    # Adding format and content to paragraph

    format = para.paragraph_format
    format.space_before = Pt(0)
    format.space_after = Pt(3)
    format.line_spacing = 1.0
    format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(main_title)
    run.font.name = 'CastleTLig'
    run.font.size = Pt(26)

    ## ----- Page Subheading

    # Creating paragraph
    para = doc.add_paragraph()

    # Adding format to paragraph

    format = para.paragraph_format
    format.space_before = Pt(0)
    format.space_after = Pt(4)
    format.line_spacing = 1.0
    format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adding page title to paragraph

    if page_description and page_description.strip():
        sub_text = page_description.upper()
    elif label_print:
        sub_text = "CONSTANT PLATE FLAW LABELS"
    else:
        sub_text = title.upper()

    run = para.add_run(sub_text)
    run.font.name = 'CastleTLig'
    run.font.size = Pt(13)


def printDocxPages(doc1, first_page, temp_list, reference_year, title=None) :
    """Render plate-fault pages into a python-docx document.

    Arranges ``temp_list`` into a grid of pages and rows, adds headings,
    inserts stamp images and descriptions into tables, and records per-page
    timing data.

    Args:
        doc1: ``python-docx`` document to populate.
        first_page: If True, skip the initial page break before the first page.
        temp_list: Plate-fault data rows to render.
        reference_year: Reference year passed through to ``printHeadings()``.

    Side Effects:
        Mutates ``doc1`` in place and appends detailed timing data via
        ``_write_detailed_timings()``.
    """
    
  
    ## ----- Table of rows and columns of plate fault images
    ##       and descriptions
    

    if label_print :
        num_faults_horizontally = 4  
        num_faults_vertically = 10
    else :
        num_faults_horizontally = 5
        num_faults_vertically = 6
              
    ## rearrage data_list to show items within rows within pages  
    
    list2 = [temp_list[i:i+num_faults_horizontally] for i in range(0, len(temp_list), num_faults_horizontally)]
    list3 = [list2[i:i+num_faults_vertically] for i in range(0, len(list2), num_faults_vertically)]
                              
    '''
    print('\nDiagnostics ------------------------------------------------')
    for pg in list3 :
        print('----- page')
        for row in pg :
            print('      ----- row')
            for cell  in row :
                print('            -----', cell[i_rdesc])
    '''
    
    print('')
    
    # Prepare collectors for detailed timing
    page_timings = []  # list of (page_index, page_duration, num_images_on_page, avg_image_time)
    total_image_time = 0.0
    total_image_count = 0

    page_index = 0

    for page_list in list3 :
        if not first_page:
            doc1.add_page_break() 
            [doc1.add_paragraph('') for _ in range(1)]
        else :
            first_page = False
            pass
        printHeadings(doc1, reference_year, title=title)

        pad_text = '\n'

        page_start = time.perf_counter()
        image_count_this_page = 0

        for row_list in page_list :
            
            ## create a table in the doc1 document with 1 row and a number of
            #    - columns equal the value of num_faults_horizontally in the referece documents
            #    - columns equal to twice the value of num_faults_horizontally in the labels documents
            #  and applies the "Table Grid" style to the labels documents
            
            if label_print :
                table = doc1.add_table(rows=1, cols=num_faults_horizontally*2, style='Table Grid')
            else :
                table = doc1.add_table(rows=1, cols=num_faults_horizontally)
            table.allow_autofit = True

            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.allow_autofit = True
            
            if label_print:
                row = table.rows[0]
                row.height = Inches(0.89)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            
            cells = table.rows[0].cells

            for index, item in enumerate(row_list) :
                img_start = time.perf_counter()
                
                if label_print : ## if label, print image and description horizontally
                    index = index * 2

                ## insert image into docx table
                
                pic_path = cleanUpPathName(images_directory)[0:-6] + '/' + item[1]
                cell = cells[index]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell_p, cell_f, cell_r = paragraph_format_run(cell)
                temp = images_directory[0:-6] + item[1]
                try:
                    with open(temp) :
                        isExist = True
                except IOError:
                    isExist = False
                if isExist :
                    if label_print :
                        cell_r._element.getparent().height = Inches(1.5)
                        cell_r.add_picture(pic_path, height=Inches(.875), width=Inches(.875))
                    else :
                        cell_r.add_picture(pic_path, height=Inches(1.05), width=Inches(1.05))
                else:
                    if label_print :
                        cell_r.font.size=Pt(4.5)
                        cell_r.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        cell_r.font.name = 'Verdana'
                        cell_r.add_text('(Image Missing)')
                    else :
                        cell_r.text = '\n\n\n\n\n\n(Image Missing)\n\n\n\n\n\n\n '
                
                ## insert plate fault description into docx table
                
                if label_print :
                    if index < num_faults_horizontally*2-1 :
                        index += 1
                
               
                cell = cells[index]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                cell_p, cell_f, cell_r = paragraph_format_run(cell)
                cell_r.text = pad_text + cleanUpPFDesc(item[9])
                
                ## loop through the runs and set font and font size
                
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Verdana'
                        font.size= Pt(4.5)
                
                ## loop through the cells and set the border color to light if labels
                
                if label_print :
                    for row in table.rows:
                        for cell in row.cells:
                            set_cell_border(
                                cell,
                                top={"sz": 2, "val": "dotted", "color": "#d8d8d8", "space": "0"},
                                bottom={"sz": 2, "val": "dotted", "color": "#d8d8d8", "space": "0"},
                                start={"sz": 2, "val": "dotted", "color": "#d8d8d8", "space": "0"},
                                end={"sz": 2, "val": "dotted", "color": "#d8d8d8", "space": "0"},
                            )
                
                # end of processing for this image+caption cell
                img_end = time.perf_counter()
                img_dur = img_end - img_start
                total_image_time += img_dur
                total_image_count += 1
                image_count_this_page += 1

        page_end = time.perf_counter()
        page_duration = page_end - page_start
        avg_image_time = (page_duration / image_count_this_page) if image_count_this_page else 0.0
        page_timings.append((page_index, page_duration, image_count_this_page, avg_image_time))
        page_index += 1

    # write detailed timings for this run
    try:
        _write_detailed_timings(page_timings, total_image_time, total_image_count)
    except Exception as e:
        print(f'Warning: failed to write detailed timings: {e}')

def processAlbum(line):
    """Handle an ``<album>`` XML opening tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``<album>``.
    """
    return

def processColumnSet(line):
    """Handle a ``<column-set>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line for a column-set element.
    """
    return

def processComment(line):
    """Handle an XML comment line.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML comment line starting with ``<!--``.
    """
    return

def processCompSet(line):
    """Handle a ``<comp-set>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line for a comp-set element.
    """
    return
def processContentItems(line):
    """Handle a ``<content-items>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line for a content-items element.
    """
    return

def processDoctype(line):
    """Handle a DOCTYPE declaration line.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line containing ``!DOCTYPE``.
    """
    return

def processItem(line):
    """Handle an ``<item>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``<item>``.
    """
    return

def processPage(line):
    """Extract page title and subtitle from a ``<page>`` XML tag.

    Parses ``title`` and ``subtitle`` attributes from the line and stores
    them in module-level ``page_title`` and ``page_description``.

    Args:
        line: Raw XML line starting with ``<page ``.

    Returns:
        list: Zero, one, or two entries of the form ``['page_title', ...]``
        and/or ``['page_subtitle', ...]``.

    Side Effects:
        Updates global ``page_title`` and ``page_description``.
    """
    global page_title, page_description
    
    ans1 = []
    ans2 = []
    ans = []
    if 'title' in line :
        page_title = line.split('title')[1].split('"')[1].split('"')[0]
        ans1 = ['page_title', page_title]
        ans.append(ans1)
    if 'subtitle' in line :
        page_description = line.split('subtitle')[1].split('"')[1].split('"')[0]
        ans2 = ['page_subtitle', page_description]
        ans.append(ans2)
    return ans

def processRowSet(line):
    """Extract a row-set description from a ``<row-set>`` XML tag.

    Parses the ``description`` attribute when present. Returns an empty
    description for bare ``<s>`` lines.

    Args:
        line: Raw XML line for a row-set element.

    Returns:
        list: ``['row_set_description', description_text]``.

    Side Effects:
        Updates global ``f_row_set_description``.
    """
    global f_row_set_description, row_set_description
    
    if line[0:3] =='<s>' :
        f_row_set_description = ''
    else :
        if 'description=' in line :
            f_row_set_description = line.split('description=')[1].split('"')[1].split('"')[0]                             
        else:
            f_row_set_description = 'ERROR'  ## This should never be the case
      
    return [ 'row_set_description', f_row_set_description]

def processS(line):
    """Parse a stamp ``<s>`` XML line into a data row.

    Extracts image path, dimensions, denomination, color, design, and catalog
    numbers from quoted attribute values in the XML line.

    Args:
        line: Raw XML line for an ``<s>`` or ``<s ...>`` element.

    Returns:
        list: Data row of the form ``['data', image, dimensions, denomination,
        color, design, cat_nums, cat1, cat2]``.
    """
        
    ## check for proper number of quote ("") pairs
    
    if line.count('"') != 10 :  ## guard against improper number of fields
        #print('\nERROR: incorrect number of fields\n', line)
        1==1
    
    ## extract image

    if line[0:3] =='<s>' :
        f_image = ''
    else :
        if 'image=' in line :
            f_image = line.split('image=')[1].split('"')[1].split('"')[0]                             
        else:
            f_image = ''  # case of <shape
     
    ## extract dimensions (clean up text in print routines if needed)
    
    f_dimensions = line.split('>')[1].split('"')[1] 
    if len(f_dimensions) > 2 :
        f_dimensions = f_dimensions.split()[0].strip() + 'x' + f_dimensions.split()[1].strip()
    else:
        f_dimensions = '???'
    
    ## extract denomination
    
    f_denomination = line.split('>')[1].split('"')[3].strip()

    ## extract color
   
    f_color = line.split('>')[1].split('"')[5].strip()
                           
    ## Extract design
    
    f_design = line.split('>')[1].split('"')[7].strip()
    
    ## Extract catalog numbers
   
    if len(line) != 8 : ## added V01 to kludge incorrect number of double quotes
        f_cat_nums = ''
    else :
        f_cat_nums = line.split('>')[1].split('"')[9].strip()
    
    ## Extract First Catalog Number
    if f_cat_nums == '' : ## case where catalog numbers is empty
        f_cat1 = ''
    elif f_cat_nums == '-' :  ## case where catalog numbers is '-'
        f_cat1 = '-'
    elif f_cat_nums.count('(') == 0 :  ## case where catalog numbers is singular
        f_cat1 = ''
    else:
        f_cat1 = f_cat_nums.split('(')[0].strip()
    
    ## Extract Second Catalog Number
   
    if f_cat_nums == '' : ## case where catalog numbers is empty
        f_cat2 = ''
    elif f_cat_nums == '-' :  ## case where catalog numbers is '-'
        f_cat2 = '-'
    elif f_cat_nums.count('(') == 0 :  ## case where catalog numbers is singular
        f_cat2 = ''
    else:
        f_cat2 = f_cat_nums.split('(')[1].split(')')[0].strip()
                
    return ['data', f_image, f_dimensions, f_denomination, f_color, f_design, f_cat_nums, f_cat1, f_cat2]

def processSet(line):
    """Extract the issue attribute from a ``<set>`` XML tag.

    Args:
        line: Raw XML line starting with ``<set``.

    Returns:
        list: ``['set', issue_text]``, where ``issue_text`` is empty when no
        ``issue`` attribute is present.
    """
    issue = ''
    if 'issue="' in line :
        issue = line.split('issue="')[1].split('"')[0]
    return ['set', issue]

def processTitlePage(line):
    """Handle a ``<title-page>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line for a title-page element.
    """
    return

def processZAlbum(line):
    """Handle a closing ``</album>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</album>``.
    """
    return

def processZColumnSet(line):
    """Handle a closing ``</column-set>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</column-set>``.
    """
    return

def processZCompSet(line):
    """Handle a closing ``</comp-set>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</comp-set>``.
    """
    return

def processZContentItems(line):
    """Handle a closing ``</content-items>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</content-items>``.
    """
    return

def processZItem(line):
    """Handle a closing ``</item>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</item>``.
    """
    return

def processZPage(line):
    """Handle a closing ``</page>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</page>``.
    """
    return

def processZRowSet(line):
    """Handle a closing ``</row-set>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</row-set>``.
    """
    return

def processZS(line):
    """Handle a closing ``</s>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</s>``.
    """
    return

def processZSet(line):
    """Handle a closing ``</set>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</set>``.
    """
    return

def processZSetTenant(line):
    """Handle a closing ``</set-tenant>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</set-tenant>``.
    """
    return

def processZTitlePage(line):
    """Handle a closing ``</title-page>`` XML tag.

    Placeholder handler invoked by ``create_data_list()``; no processing is
    performed.

    Args:
        line: Raw XML line starting with ``</title-page>``.
    """
    return

def removeDenominationColor(data_list):
    """Keep only data rows with empty denomination, color, and cell description.

    Filters plate-fault ``data`` entries to those without denomination, color,
    or cell-description values, which identifies rows suitable for PF
    reference output.

    Args:
        data_list: Parsed album data from earlier pipeline steps.

    Returns:
        list: Filtered data rows with empty denomination, color, and
        cell-description fields.
    """
    temp = []
    for entry in data_list :
        if entry[0] == 'data' and entry[i_denomination] == '' and entry[i_color] == '' and entry[i_cell_desc]  == '' :
            temp.append(entry)
    return temp

def remove_duplicates_from_data_list(data_list, i_image, i_rdesc):
    """Remove duplicate plate-fault rows by image and description.

    Keeps the first occurrence of each unique ``(image, description)`` pair
    and preserves original order.

    Args:
        data_list: Plate-fault data rows to deduplicate.
        i_image: Index of the image field within each data row.
        i_rdesc: Index of the description field within each data row.

    Returns:
        list: Deduplicated data rows.

    Side Effects:
        Prints before-and-after plate-fault counts to the console.
    """
    
    print(f'Initial plate flaw count: {len(data_list)}') 
    
    # Set to store unique combinations of image and description
    seen = set()
    # List to store non-duplicate items while preserving order
    result = []

    # Iterate through each sublist in the data_list
    for item in data_list:
        # Create a tuple of the image and description values
        # This tuple will serve as our unique identifier
        identifier = (item[i_image], item[i_rdesc])

        # Check if this identifier has been seen before
        if identifier not in seen:
            # If not seen, add to the set of seen identifiers
            seen.add(identifier)
            # Add the entire sublist to the result
            result.append(item)

    print(f'Processed plate flaw count: {len(result)}') 

    return result

def set_cell_border(cell, **kwargs):
    """Apply Word table-cell border styling via OOXML.

    Sets one or more cell edges using python-docx low-level XML. Supported
    edge keys are ``start``, ``top``, ``end``, ``bottom``, ``insideH``, and
    ``insideV``. Each edge accepts border attributes such as ``sz``, ``val``,
    ``color``, ``space``, and ``shadow``.

    Args:
        cell: ``python-docx`` table cell to modify.
        **kwargs: Border definitions keyed by edge name. Example::

            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            )

    Side Effects:
        Mutates the cell's underlying OOXML border elements in place.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))

def set_repeat_table_header(row):
    """Mark a table row to repeat as a header on each page.

    Args:
        row: ``python-docx`` table row to mark as a repeating header.

    Returns:
        The same row object, after OOXML header flags are applied.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def splitCombinedXML(combined_xml):
    parts = combined_xml.split('<')
    return ['<' + part.strip() for part in parts[1:]]

def writeDocxPfRefDoc(doc):
    """Save the PF reference document as a DOCX file.

    Writes ``doc`` to the path derived from global ``output_file_wo_extension``.
    If the target file is locked (e.g. open in Word), retries with a ``(1)``
    suffix.

    Args:
        doc: Completed ``python-docx`` document to save.

    Side Effects:
        Updates global ``docx_file`` and possibly ``output_file_wo_extension``.
        Prints status or error messages to the console.
    """
    
    global docx_file, output_file_wo_extension
    
    docx_file = output_file_wo_extension + '.docx'
    
    print(f'Creating PF Reference Document (Docx) at {docx_file}')    
    
    try :
        doc.save (docx_file) 
    except :
        print ('ERROR — Word Document Is Open')
        output_file_wo_extension = output_file_wo_extension + ' (1)'
        docx_file = output_file_wo_extension + '.docx'
        doc.save (docx_file)
    #word.Quit()
    
def writePdfPfRefDoc():
    """Convert the saved DOCX file to PDF using Microsoft Word COM.

    Opens the DOCX at global ``docx_file`` via Word automation and saves it
    as PDF. Uses alternate ``(1)`` filenames when source or destination files
    are locked or unavailable.

    Returns:
        str | None: Path to the saved PDF on success, or ``None`` on failure.

    Side Effects:
        Launches and quits a Word application instance. Prints status and
        error messages to the console.
    """
    import os
    from pathlib import Path
    
    global docx_file, output_file_wo_extension
    
    try:
        # Normalize paths using pathlib
        output_base = Path(output_file_wo_extension).resolve()
        
        if label_print:
            dst = output_base.parent / f"{output_base.name} Labels.pdf"
            print(f'Creating PF Labels Document (PDF) at {dst}')
        else:
            dst = output_base.with_suffix('.pdf')
            print(f'Creating PF Reference Document (PDF) at {dst}')
        
        # Ensure source path is absolute and exists
        src = Path(docx_file).resolve()
        if not src.exists():
            raise FileNotFoundError(f"Source file not found: {src}")

        # Initialize Word
        if not HAS_WIN32COM or win32com is None:
            print("ERROR — Microsoft Word COM automation is unavailable on this platform (Word processor mode requires Windows and MS Word).")
            return None

        try:
            word = win32com.client.Dispatch("Word.Application")
        except Exception as e:
            err_msg = str(e)
            if "-2147221005" in err_msg or "Invalid class string" in err_msg or "coinitialize" in err_msg.lower():
                print("ERROR — Microsoft Word COM automation is unavailable on this system (Word processor mode requires Microsoft Word).")
            else:
                print(f"ERROR — Failed to initialize Microsoft Word application: {err_msg}")
            return None

        word.Visible = False
        wdFormatPDF = 17
        
        try:
            # Ensure directory exists
            dst.parent.mkdir(parents=True, exist_ok=True)
            
            # Convert paths to string for Word COM API
            src_str = str(src)
            dst_str = str(dst)
            
            try:
                print(f"Attempting to open: {src_str}")
                doc = word.Documents.Open(src_str)
                print("Document opened successfully")
            except Exception as e:
                print(f'ERROR — Could not open Word document: {str(e)}')
                # Try with (1) suffix if the file is open
                alt_src = output_base.parent / f"{output_base.name} (1).docx"
                print(f'Trying with alternative filename: {alt_src}')
                alt_src_str = str(alt_src)
                if alt_src.exists():
                    doc = word.Documents.Open(alt_src_str)
                    src_str = alt_src_str
                else:
                    raise FileNotFoundError(f"Alternative file not found: {alt_src}")
            
            try:
                print(f"Attempting to save as PDF: {dst_str}")
                doc.SaveAs(dst_str, FileFormat=wdFormatPDF)
                print("PDF saved successfully")
            except Exception as e:
                print(f'ERROR — Could not save PDF: {str(e)}')
                # Try with (1) suffix if the file is open
                alt_dst = output_base.parent / f"{output_base.name} (1).pdf"
                print(f'Trying with alternative filename: {alt_dst}')
                dst_str = str(alt_dst)
                doc.SaveAs(dst_str, FileFormat=wdFormatPDF)
            
            doc.Close(False)  # Don't save changes
            print(f"Generated PDF: {os.path.abspath(str(dst_str))}")
            return dst_str
            
        except Exception as e:
            print(f'Error in PDF conversion: {str(e)}')
            return None
            
        finally:
            word.Quit()
            
    except Exception as e:
        print(f'Error in writePdfPfRefDoc: {str(e)}')
        return None


def _write_timings_log(timings, note=None):
    """Append pipeline timing measurements to a baseline log file.

    Writes one timestamped entry to ``baseline_timings.txt`` in
    ``output_directory``, including optional context from globals.

    Args:
        timings: Mapping of phase names to elapsed seconds or other values.
        note: Optional note appended to the log entry, such as ``'no_data_found'``.

    Side Effects:
        Creates ``output_directory`` if needed and appends to the log file.
        Prints a warning if logging fails.
    """
    try:
        try:
            out_base = output_file_wo_extension
        except NameError:
            out_base = None
        try:
            title = constant_title
        except NameError:
            title = 'N/A'

        logs_dir = os.path.join(output_directory, 'logs')
        os.makedirs(logs_dir, exist_ok=True)
        log_file = os.path.join(logs_dir, 'baseline_timings.txt')
        with open(log_file, 'a', encoding='utf-8') as fh:
            fh.write('--- TIMING: ' + time.strftime("%Y-%m-%d %H:%M:%S") + '\n')
            fh.write(f'Title: {title}\n')
            if out_base:
                fh.write(f'Output base: {out_base}\n')
            if note:
                fh.write(f'NOTE: {note}\n')
            for k in sorted(timings.keys()):
                try:
                    fh.write(f'{k}: {timings[k]:.4f} seconds\n')
                except Exception:
                    fh.write(f'{k}: {timings[k]}\n')
            fh.write('\n')
    except Exception as e:
        # Best-effort logging; do not raise on logging failures
        print(f'Warning: could not write timings log: {e}')


def _write_detailed_timings(page_timings, total_image_time, total_image_count, note=None):
    """Append per-page image timing details to a log file.

    Writes aggregated and per-page DOCX rendering timings to
    ``detailed_timings.txt`` in ``output_directory/logs``.

    Args:
        page_timings: List of tuples ``(page_index, page_duration_s,
            num_images, avg_image_time_on_page_s)``.
        total_image_time: Sum of per-image processing durations, in seconds.
        total_image_count: Total number of images processed.
        note: Optional note appended to the log entry.

    Side Effects:
        Creates ``output_directory/logs`` if needed and appends to the log file.
        Prints a warning if logging fails.
    """
    try:
        logs_dir = os.path.join(output_directory, 'logs')
        os.makedirs(logs_dir, exist_ok=True)
        log_file = os.path.join(logs_dir, 'detailed_timings.txt')
        with open(log_file, 'a', encoding='utf-8') as fh:
            fh.write('--- DETAILED TIMING: ' + time.strftime("%Y-%m-%d %H:%M:%S") + '\n')
            try:
                fh.write(f'Title: {constant_title}\n')
            except Exception:
                fh.write('Title: N/A\n')
            if note:
                fh.write(f'NOTE: {note}\n')
            fh.write(f'Total images processed: {total_image_count}\n')
            fh.write(f'Total image processing time (sum of per-image durations): {total_image_time:.4f} seconds\n')
            if total_image_count:
                fh.write(f'Average image processing time (global): {total_image_time/total_image_count:.6f} seconds\n')
            fh.write('\n')
            fh.write('Per-page timings (page_index, page_duration_s, num_images, avg_image_time_on_page_s)\n')
            for p in page_timings:
                fh.write(f'{p[0]}: {p[1]:.4f}s, images={p[2]}, avg_image={p[3]:.6f}s\n')
            fh.write('\n')
    except Exception as e:
        print(f'Warning: could not write detailed timings file: {e}')





def get_selection():
    """Resolve configuration parameters following 3-tier precedence hierarchy:
    Tier 1 (Defaults) < Tier 2 (mapping.json) < Tier 3 (CLI args).

    Returns:
        tuple[str, str, str, str, str]: ``(selection_key, processor, input_dir, output_dir, images_dir)``
    """
    import argparse

    parser = argparse.ArgumentParser(description='Generate PF Reference from XML files')
    parser.add_argument('--selection', '-s', type=str.lower, 
                       choices=selection_map.keys(),
                       help='Selection from available options')
    parser.add_argument('--processor', '-p', type=str.lower, choices=['word', 'reportlab'], default=None, help='PDF processor to use')
    parser.add_argument('--input-dir', '-i', type=str, default=None, help='Base directory for XML input files')
    parser.add_argument('--output-dir', '-o', type=str, default=None, help='Directory for generated PDF outputs')
    parser.add_argument('--images-dir', '-img', type=str, default=None, help='Directory for stamp image assets')
    parser.add_argument('--image-type', type=str, choices=['png', 'jpeg', 'jpg'], default=None, help='Target image format')
    parser.add_argument('--image-quality', '-q', type=int, default=None, help='Target image quality (1-100)')
    parser.add_argument('--debug', '-d', action='store_true', help='Enable debug logging')

    # Only parse known arguments to avoid conflicts with other modules
    args, _ = parser.parse_known_args()

    log_level = logging.DEBUG if getattr(args, 'debug', False) else logging.WARNING
    logging.basicConfig(level=log_level, format='%(levelname)s: %(message)s', force=True)

    resolved = _resolve_config_params(cli_args=args)

    selection = resolved['selection']
    proc = resolved['processor']
    input_dir = resolved['input-dir']
    output_dir = resolved['output-dir']
    images_dir = resolved['images-dir']

    # Console logging if --debug is enabled
    logging.debug(f"get_selection: parsed selection={selection}, processor={proc}, input_dir={input_dir}, output_dir={output_dir}, images_dir={images_dir}")

    if selection and selection in selection_map:
        return selection, proc, input_dir, output_dir, images_dir

    # If no selection specified via CLI or mapping.json, prompt interactively
    while True:
        prompt_selection = input(f"Enter your selection ({', '.join(selection_map.keys())}): ").lower()
        if prompt_selection in selection_map:
            return prompt_selection, proc, input_dir, output_dir, images_dir
        print(f"Invalid selection: {prompt_selection}")

def main():
    """Run the PF reference generation workflow for the chosen album.

    Resolves user selection, builds input/output paths, generates constant
    plate flaw pages, then constant overprint flaw pages, and prints elapsed
    runtime.

    Side Effects:
        Updates globals ``constant_title``, ``output_file``, and
        ``output_file_wo_extension``. Calls ``createPFRefAlbumPages()`` twice
        and prints start/end timestamps to the console.
    """
    global constant_title, output_file, output_file_wo_extension, input_directory, output_directory, images_directory
    
    # Get selection and CLI directories from command line or interactive input
    selection, proc, cli_input_dir, cli_output_dir, cli_images_dir = get_selection()

    if cli_input_dir:
        input_directory = cli_input_dir
        images_directory = os.path.join(input_directory, 'images')
        # Re-build selection_map lists with the updated input_directory
        _rebuild_selection_map()

    if cli_output_dir:
        output_directory = cli_output_dir

    if cli_images_dir:
        images_directory = cli_images_dir

    # Process the selection
    input_file_list, output_filename = selection_map[selection]
    output_file = os.path.join(output_directory, output_filename)
    output_file_wo_extension = output_file.split('.')[0]
    
    start_time = time.time()  # Get the current time in seconds since the epoch
    current_time_str = time.strftime("%Y-%m-%d %H:%M:%S")  # Format the current time as a string
    print(f"Start Time: {current_time_str}")  # Print the current time to the console


    ## Create consolidated PF reference pages (Constant Plate Flaws and Constant Overprint Flaws)
    constant_title = "Constant Plate Flaws"
    createPFRefAlbumPages(input_file_list, output_file, images_directory, processor=proc)
    
    current_time_str_again = time.strftime("%Y-%m-%d %H:%M:%S")  # Get the current time again
    print(f"End time: {current_time_str_again}")  # Print the current time after the greeting
    end_time = time.time()  # Record the end time
    elapsed_time = end_time - start_time  # Calculate the elapsed time in seconds
    if elapsed_time >= 60:  # Check if elapsed time is 60 seconds or more
        minutes = int(elapsed_time // 60)  # Calculate full minutes
        seconds = elapsed_time % 60  # Calculate remaining seconds
        print(f"Elapsed Time: {minutes} minutes and {seconds:.2f} seconds")  # Print in minutes and seconds
    else:  # If elapsed time is less than 60 seconds
        print(f"Elapsed Time: {elapsed_time:.2f} seconds")  # Print formatted to two decimal places


def _load_selection_mapping(input_dir=None):
    """Load selection definitions from selection-mapping.json.
    
    Returns:
        dict: {selection_key: (file_list, output_filename)}
    """
    if input_dir is None:
        input_dir = input_directory

    json_path = Path('selection-mapping.json')
    if not json_path.exists():
        json_path = Path(__file__).parent / 'selection-mapping.json'

    raw_mapping = {}
    if json_path.exists():
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                raw_mapping = json.load(f)
        except Exception as e:
            print(f"Warning: Failed to load selection-mapping from {json_path}: {e}")

    s_map = {}
    for key, info in raw_mapping.items():
        if isinstance(info, dict) and 'files' in info and 'output' in info:
            file_list = create_file_list(input_dir, info['files'])
            s_map[key.strip().lower()] = (file_list, info['output'])

    return s_map


def _rebuild_selection_map():
    global selection_map
    selection_map = _load_selection_mapping(input_directory)


# Map input selections to their respective file lists and output filenames
selection_map = _load_selection_mapping(input_directory)


if __name__ == "__main__":
    main()

